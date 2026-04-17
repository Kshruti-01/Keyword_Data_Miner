"""
Word Document Generator for Keyword Mining Results
Creates formatted Word documents from keyword extraction results.
"""

import os
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("python-docx not installed. Word generation disabled.")
    print(" Install with: pip install python-docx")


class WordReportGenerator:
    """
    Generates formatted Word documents from keyword mining results.
    """
    
    def __init__(self):
        self.document = None
    
    def create_report(self, results, keywords, total_emails, emails_with_keywords, 
                      total_keywords_found, output_path):
        """Create a comprehensive Word report."""
        if not DOCX_AVAILABLE:
            return None
        
        self.document = Document()
        self._add_header()
        self._add_summary_section(total_emails, emails_with_keywords, total_keywords_found, keywords)
        self._add_keyword_frequency_table(results, keywords)
        self._add_detailed_results(results)
        self._add_footer()
        self.document.save(output_path)
        print(f" Word report saved to: {output_path}")
        return output_path
    
    def _add_header(self):
        """Add title and header to the document."""
        title = self.document.add_heading('Email Keyword Mining Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        date_para = self.document.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.add_run(f'Generated on: {datetime.now().strftime("%B %d, %Y at %H:%M:%S")}')
        self.document.add_paragraph('_' * 50)
    
    def _add_summary_section(self, total_emails, emails_with_keywords, total_keywords_found, keywords):
        """Add executive summary section."""
        self.document.add_heading('Executive Summary', level=1)
        
        table = self.document.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        
        success_rate = (emails_with_keywords / total_emails * 100) if total_emails > 0 else 0
        
        cells = [
            ("Total Emails Processed", str(total_emails)),
            ("Emails with Keywords Found", str(emails_with_keywords)),
            ("Success Rate", f"{success_rate:.1f}%"),
            ("Total Keywords Found", str(total_keywords_found)),
            ("Keywords Searched", ", ".join(keywords[:15]) + ("..." if len(keywords) > 15 else ""))
        ]
        
        for i, (label, value) in enumerate(cells):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            row.cells[0].paragraphs[0].runs[0].bold = True
        
        self.document.add_paragraph()
    
    def _add_keyword_frequency_table(self, results, keywords):
        """Add keyword frequency analysis table."""
        self.document.add_heading('Keyword Frequency Analysis', level=1)
        
        keyword_freq = {kw: 0 for kw in keywords}
        for result in results:
            for keyword in result.get('keywords', {}):
                if keyword in keyword_freq:
                    keyword_freq[keyword] += 1
        
        keyword_freq = {k: v for k, v in keyword_freq.items() if v > 0}
        sorted_keywords = sorted(keyword_freq.items(), key=lambda x: x[1], reverse=True)
        
        if sorted_keywords:
            table = self.document.add_table(rows=len(sorted_keywords) + 1, cols=3)
            table.style = 'Light Grid Accent 1'
            
            headers = table.rows[0].cells
            headers[0].text = "Rank"
            headers[1].text = "Keyword"
            headers[2].text = "Frequency (Emails)"
            
            for cell in headers:
                cell.paragraphs[0].runs[0].bold = True
            
            for i, (keyword, count) in enumerate(sorted_keywords, 1):
                row = table.rows[i]
                row.cells[0].text = str(i)
                row.cells[1].text = keyword
                row.cells[2].text = str(count)
                
                if count >= 5:
                    row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 102, 204)
                    row.cells[1].paragraphs[0].runs[0].bold = True
            
            self.document.add_paragraph()
        else:
            self.document.add_paragraph("No keywords found in any emails.")
    
    def _add_detailed_results(self, results):
        """Add detailed results for each email."""
        self.document.add_heading('Detailed Email Analysis', level=1)
        
        for i, result in enumerate(results[:20], 1):
            email_meta = result.get('email_metadata', {})
            keywords = result.get('keywords', {})
            
            subject = email_meta.get("subject", "No Subject")[:80]
            self.document.add_heading(f'Email {i}: {subject}', level=2)
            
            metadata_table = self.document.add_table(rows=3, cols=2)
            metadata_table.style = 'Light Shading'
            
            metadata = [
                ("From", email_meta.get("sender", "Unknown")),
                ("Date", email_meta.get("date", "Unknown")),
                ("Keywords Found", str(len(keywords)))
            ]
            
            for j, (label, value) in enumerate(metadata):
                row = metadata_table.rows[j]
                row.cells[0].text = label
                row.cells[1].text = value
                row.cells[0].paragraphs[0].runs[0].bold = True
            
            if keywords:
                self.document.add_paragraph()
                self.document.add_heading('Keywords Found:', level=3)
                
                for keyword, data in keywords.items():
                    confidence = int(data.get('confidence', 0) * 100)
                    occurrences = data.get('occurrences', 0)
                    summary = data.get('summary', 'No summary available')
                    
                    p = self.document.add_paragraph()
                    run = p.add_run(f"{keyword.upper()}")
                    run.bold = True
                    run.font.size = Pt(12)
                    
                    p.add_run(f"\n   Confidence: {confidence}% | Occurrences: {occurrences}")
                    
                    if summary:
                        p.add_run(f"\n   Summary: {summary[:200]}...")
                    
                    contexts = data.get('contexts', [])
                    if contexts:
                        ctx = contexts[0]
                        before = ctx.get('before', '')[-50:]
                        kw = ctx.get('keyword', '')
                        after = ctx.get('after', '')[:50]
                        p.add_run(f"\n   Example: ...{before} {kw} {after}...")
                    
                    p.paragraph_format.space_after = Pt(12)
            else:
                self.document.add_paragraph("No keywords found in this email.")
            
            self.document.add_paragraph('_' * 80)
    
    def _add_footer(self):
        """Add footer to the document."""
        self.document.add_page_break()
        footer_para = self.document.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.add_run(f'Report generated by Email Keyword Miner v1.0')
        footer_para.font.size = Pt(10)
        footer_para.font.italic = True


def create_focused_report(summary_data, search_keyword, subject_filter, output_path):
    """
    Create a focused Word report for the searched keyword within specific emails.
    
    Args:
        summary_data: Dictionary containing all extraction results
        search_keyword: The keyword that was searched for
        subject_filter: The email subject filter used to find emails
        output_path: Path where to save the Word document
    
    Returns:
        Path to the saved Word document, or None if failed
    """
    if not DOCX_AVAILABLE:
        print("python-docx not installed. Word generation disabled.")
        return None
    
    try:
        doc = Document()
        
        # Title Section
        title = doc.add_heading('Focused Keyword Search Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run(f'Email Subject Filter: "{subject_filter}"')
        subtitle_run.bold = True
        subtitle.add_run(f'\nKeyword Searched: "{search_keyword}"')
        
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.add_run(f'Generated on: {datetime.now().strftime("%B %d, %Y at %H:%M:%S")}')
        
        doc.add_paragraph('_' * 50)
        
        # Search Statistics Section
        doc.add_heading('Search Statistics', level=1)
        
        stats_table = doc.add_table(rows=5, cols=2)
        stats_table.style = 'Light Grid Accent 1'
        
        for row in stats_table.rows:
            row.cells[0].paragraphs[0].runs[0].bold = True
        
        stats = [
            ("Email Subject Filter", subject_filter),
            ("Keyword Searched", search_keyword),
            ("Emails with Keyword", str(summary_data.get('emails_with_keyword', 0))),
            ("Total Occurrences", str(summary_data.get('total_occurrences', 0))),
            ("Success Rate", f"{summary_data.get('success_rate', 0):.1f}%")
        ]
        
        for i, (label, value) in enumerate(stats):
            stats_table.rows[i].cells[0].text = label
            stats_table.rows[i].cells[1].text = value
        
        doc.add_paragraph()
        
        # Context Examples Section
        contexts = summary_data.get('contexts', [])
        if contexts:
            doc.add_heading('Context Examples', level=1)
            doc.add_paragraph('Here are examples of how the keyword appears in the emails:')
            
            for i, ctx in enumerate(contexts[:5], 1):
                p = doc.add_paragraph()
                p.add_run(f"Example {i}:").bold = True
                
                full_context = ctx.get('full_context', '')
                if full_context:
                    keyword_lower = search_keyword.lower()
                    text_lower = full_context.lower()
                    pos = text_lower.find(keyword_lower)
                    
                    if pos != -1:
                        before = full_context[max(0, pos-50):pos]
                        keyword_text = full_context[pos:pos+len(search_keyword)]
                        after = full_context[pos+len(search_keyword):pos+100]
                        
                        p = doc.add_paragraph()
                        p.add_run(f"...{before}")
                        highlight_run = p.add_run(keyword_text)
                        highlight_run.bold = True
                        highlight_run.font.color.rgb = RGBColor(255, 0, 0)
                        p.add_run(f"{after}...")
                    else:
                        p = doc.add_paragraph(f"...{full_context[:200]}...")
                
                p.paragraph_format.space_after = Pt(12)
        
        # Key Summaries Section
        summaries = summary_data.get('summaries', [])
        if summaries:
            doc.add_heading('Key Insights', level=1)
            doc.add_paragraph('The following insights were extracted from the emails:')
            
            for i, summary in enumerate(summaries[:3], 1):
                p = doc.add_paragraph()
                p.add_run(f"Insight {i}:").bold = True
                p.add_run(f"\n   {summary[:400]}...")
                p.paragraph_format.space_after = Pt(12)
        
        # Emails with Keyword Section
        doc.add_heading(f'Emails Containing "{search_keyword}"', level=1)
        
        results = summary_data.get('detailed_results', [])
        email_count = 0
        
        for result in results:
            keywords = result.get('keywords', {})
            if search_keyword.lower() in keywords:
                email_count += 1
                email_meta = result.get('email_metadata', {})
                keyword_data = keywords[search_keyword.lower()]
                confidence = int(keyword_data.get('confidence', 0) * 100)
                occurrences = keyword_data.get('occurrences', 0)
                
                doc.add_heading(f'Email {email_count}: {email_meta.get("subject", "No Subject")[:70]}', level=2)
                
                meta_table = doc.add_table(rows=4, cols=2)
                meta_table.style = 'Light Shading'
                
                meta_data = [
                    ("From", email_meta.get('sender', 'Unknown')),
                    ("Date", email_meta.get('date', 'Unknown')),
                    (f'"{search_keyword}" Occurrences', str(occurrences)),
                    ("Confidence Score", f"{confidence}%")
                ]
                
                for j, (label, value) in enumerate(meta_data):
                    meta_table.rows[j].cells[0].text = label
                    meta_table.rows[j].cells[1].text = value
                    meta_table.rows[j].cells[0].paragraphs[0].runs[0].bold = True
                
                if keyword_data.get('summary'):
                    doc.add_paragraph()
                    doc.add_heading('Summary:', level=3)
                    doc.add_paragraph(keyword_data['summary'][:400])
                
                contexts = keyword_data.get('contexts', [])
                if contexts:
                    doc.add_heading('Context Examples:', level=3)
                    for ctx_idx, ctx in enumerate(contexts[:2], 1):
                        full_context = ctx.get('full_context', '')
                        if full_context:
                            keyword_lower = search_keyword.lower()
                            text_lower = full_context.lower()
                            pos = text_lower.find(keyword_lower)
                            
                            if pos != -1:
                                before = full_context[max(0, pos-60):pos]
                                keyword_text = full_context[pos:pos+len(search_keyword)]
                                after = full_context[pos+len(search_keyword):pos+80]
                                
                                p = doc.add_paragraph()
                                p.add_run(f"Context {ctx_idx}: ...{before}")
                                highlight_run = p.add_run(keyword_text)
                                highlight_run.bold = True
                                highlight_run.font.color.rgb = RGBColor(255, 0, 0)
                                p.add_run(f"{after}...")
                            else:
                                doc.add_paragraph(f"Context {ctx_idx}: ...{full_context[:150]}...")
                
                doc.add_paragraph('_' * 60)
        
        if email_count == 0:
            doc.add_paragraph(f"No emails were found containing the keyword '{search_keyword}'.")
            doc.add_paragraph("This could be because:")
            doc.add_paragraph("  The keyword doesn't appear in any of the selected emails")
            doc.add_paragraph("  The keyword was spelled differently")
            doc.add_paragraph("   The keyword appears only in a different context")
        
        # Footer
        doc.add_page_break()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run(f'Report generated by Email Keyword Miner v2.0')
        footer_run.italic = True
        footer.font.size = Pt(10)
        
        footer.add_run(f'\nGenerated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        footer.font.size = Pt(8)
        
        doc.save(output_path)
        print(f"Focused Word report saved to: {output_path}")
        return output_path
        
    except Exception as e:
        print(f"Error creating focused Word report: {e}")
        return None


def create_quick_report(extraction_results, keywords, output_path):
    """
    Quick function to create a report from extraction results.
    
    Args:
        extraction_results: List of extraction results from DataMiner
        keywords: List of keywords that were searched
        output_path: Path to save the Word document
    
    Returns:
        Path to the saved Word document, or None if failed
    """
    generator = WordReportGenerator()
    
    total_emails = len(extraction_results)
    emails_with_keywords = sum(1 for r in extraction_results if r.get('keywords'))
    total_keywords_found = sum(len(r.get('keywords', {})) for r in extraction_results)
    
    return generator.create_report(
        results=extraction_results,
        keywords=keywords,
        total_emails=total_emails,
        emails_with_keywords=emails_with_keywords,
        total_keywords_found=total_keywords_found,
        output_path=output_path
                    )
